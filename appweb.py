import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

# إعداد الصفحة
st.set_page_config(page_title="نظام مدرسة الإمام النووي", layout="centered")
st.title("📝 مدرسة الإمام النووي")
st.subheader("تصدير النماذج (الاسم والصف) الموحدة")

uploaded_excel = st.file_uploader("ارفع ملف الطلاب (Excel)", type=["xlsx"])

if uploaded_excel:
    try:
        df = pd.read_excel(uploaded_excel)
        cols = df.columns.tolist()
        
        # اختيار الأعمدة
        col1, col2 = st.columns(2)
        with col1:
            name_col = st.selectbox("اختر عمود الأسماء:", cols)
        with col2:
            class_col = st.selectbox("اختر عمود الصفوف:", cols)
        
        # دمج العرض
        df['display'] = df[name_col].astype(str) + " - " + df[class_col].astype(str)
        selected_display = st.multiselect("اختر الطلاب:", df['display'].tolist())
        reason = st.text_input("سبب النموذج (التاريخ أو السبب):")

        if st.button("🚀 إنشاء الملف المدمج وتحميله"):
            if not selected_display:
                st.error("الرجاء اختيار طلاب أولاً")
            else:
                combined_doc = Document()
                # جلب البيانات المختارة فقط
                selected_df = df[df['display'].isin(selected_display)]
                
                for index, (idx, row) in enumerate(selected_df.iterrows()):
                    # فتح نسخة جديدة من القالب لكل طالب
                    template = Document("template.docx")
                    
                    # استبدال شامل (في الفقرات والجداول)
                    def perform_replace(doc_obj):
                        # استبدال في الفقرات العادية
                        for p in doc_obj.paragraphs:
                            if '[A]' in p.text: p.text = p.text.replace('[A]', str(row[name_col]))
                            if '[B]' in p.text: p.text = p.text.replace('[B]', str(row[class_col]))
                            if '[T]' in p.text: p.text = p.text.replace('[T]', reason)
                        
                        # استبدال في الجداول (إذا كانت الرموز داخل جدول)
                        for table in doc_obj.tables:
                            for r in table.rows:
                                for cell in r.cells:
                                    for paragraph in cell.paragraphs:
                                        if '[A]' in paragraph.text: paragraph.text = paragraph.text.replace('[A]', str(row[name_col]))
                                        if '[B]' in paragraph.text: paragraph.text = paragraph.text.replace('[B]', str(row[class_col]))
                                        if '[T]' in paragraph.text: paragraph.text = paragraph.text.replace('[T]', reason)

                    perform_replace(template)
                    
                    # إضافة محتوى القالب المعدل للمستند الكبير
                    for element in template.element.body:
                        combined_doc.element.body.append(element)
                    
                    # فاصل صفحات بين الطلاب
                    if index < len(selected_df) - 1:
                        combined_doc.add_page_break()

                # حفظ وتحميل
                target = BytesIO()
                combined_doc.save(target)
                target.seek(0)

                st.success(f"✅ تم دمج {len(selected_display)} نماذج بنجاح!")
                st.download_button(
                    label="📥 تحميل ملف النماذج الموحد",
                    data=target,
                    file_name="نماذج_الطلاب_الموحدة.docx"
                )
    except Exception as e:
        st.error(f"حدث خطأ: {e}")